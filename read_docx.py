import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Read paragraphs
    for para in doc.paragraphs:
        if para.text:
            full_text.append(para.text)
    
    # Read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text:
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

if __name__ == "__main__":
    file_path = "My_Voice_Clone_Agent_SRS.docx"
    try:
        content = read_docx(file_path)
        print(content)
    except Exception as e:
        print(f"Error reading the document: {e}") 